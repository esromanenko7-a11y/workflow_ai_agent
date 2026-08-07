import argparse
import sys
from pathlib import Path
from textwrap import dedent

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document as DocxDocument  # noqa: E402

try:
    import pymupdf as fitz  # noqa: E402
except ModuleNotFoundError:
    import fitz  # type: ignore  # noqa: E402


CATEGORIES = {
    "validation_rules": [
        "Проверка обязательных технических полей",
        "Проверка соответствия полей meta и data",
        "Проверка отсутствия лишних полей в data",
        "Проверка допустимых типов данных",
        "Проверка критичности ошибок пакета",
        "Проверка полноты описания атрибутов",
        "Проверка статусов обработки пакета",
        "Проверка правил допуска к следующему этапу",
        "Проверка кодов ошибок валидации",
        "Проверка предупреждений и блокирующих ошибок",
        "Проверка результата автоматического контроля",
        "Проверка связи бизнес-правил и технических ошибок",
    ],
    "metadata_requirements": [
        "Структура meta-файла",
        "Формат meta-файла",
        "Описание обязательных атрибутов",
        "Описание ключевых полей",
        "Версионирование meta-файла",
        "Правила заполнения описаний колонок",
        "Требования к nullable-признаку",
        "Описание первичных ключей",
        "Описание уникальных ключей",
        "Требования к стратегии загрузки",
        "Проверка порядка колонок",
        "Правила изменения metadata между версиями",
    ],
    "data_file_requirements": [
        "Структура data-файла",
        "Правила экранирования кавычек",
        "Проверка разделителей в data-файле",
        "Проверка кодировки файла",
        "Проверка пустых значений",
        "Проверка формата дат",
        "Проверка числовых значений",
        "Проверка обязательных колонок",
        "Проверка лишних колонок",
        "Проверка строк с ошибками парсинга",
        "Правила обработки больших файлов",
        "Проверка согласованности data и meta",
    ],
    "package_processing": [
        "Распаковка пакета данных",
        "Проверка структуры архива",
        "Проверка именования файлов",
        "Проверка наличия всех файлов пакета",
        "Проверка контрольных статусов",
        "Правила перехода пакета между этапами",
        "Обработка частично загруженного пакета",
        "Повторная обработка пакета",
        "Эскалация ошибок обработки",
        "Логирование результатов проверки",
        "Формирование итогового статуса пакета",
        "Передача пакета в следующий контур",
    ],
    "support_cases": [
        "Кейс: отсутствует обязательное поле",
        "Кейс: поле есть в data, но отсутствует в meta",
        "Кейс: неверный тип данных",
        "Кейс: нарушено именование файла",
        "Кейс: архив не распаковывается",
        "Кейс: предупреждение без блокировки",
        "Кейс: блокирующая ошибка",
        "Кейс: пользователь просит объяснить статус",
        "Кейс: повторная отправка исправленного пакета",
        "Кейс: конфликт версий meta-файла",
        "Кейс: неизвестная ошибка проверки",
        "Кейс: пакет можно передавать дальше",
    ],
}


def build_content(category: str, title: str, index: int) -> str:
    rule_code = (
        title.upper()
        .replace(" ", "_")
        .replace(":", "")
        .replace("-", "_")
    )

    return dedent(
        f"""
        # {title}

        Категория документа: {category}
        Код правила или сценария: {rule_code}
        Версия документа: v1.{index % 3}

        ## Назначение

        Документ описывает правило или сценарий проверки пакета данных.
        Ассистент должен использовать этот материал для объяснения ошибок,
        предупреждений и статусов обработки пакета.

        ## Условия проверки

        Проверка применяется к пакету данных на этапе автоматического контроля.
        На вход поступают meta-файл, data-файл, технические статусы обработки
        и список найденных ошибок.

        ## Логика анализа

        Если проверка относится к критичным ошибкам, пакет нельзя передавать
        дальше в обработку до исправления причины. Если найдены только
        предупреждения, пакет может быть передан дальше, но пользователь должен
        получить понятное пояснение.

        ## Рекомендация пользователю

        Необходимо проверить исходный файл, сопоставить ошибку с описанием
        правила и исправить данные или metadata. После исправления пакет нужно
        отправить на повторную проверку.

        ## Пример ответа ассистента

        Обнаружена проблема по правилу {rule_code}. Проверьте файл пакета,
        исправьте несоответствие и повторите загрузку. Если ошибка сохраняется,
        передайте пакет на эскалацию с приложением лога проверки.
        """
    ).strip()


def write_text_file(path: Path, content: str) -> None:
    path.write_text(content, encoding="utf-8")


def write_html_file(path: Path, title: str, content: str) -> None:
    html = f"""<!doctype html>
<html lang="ru">
<head>
  <meta charset="utf-8">
  <title>{title}</title>
</head>
<body>
  <article>
    {content.replace("\n", "<br>\n")}
  </article>
</body>
</html>
"""
    path.write_text(html, encoding="utf-8")


def write_docx_file(path: Path, title: str, content: str) -> None:
    document = DocxDocument()
    document.core_properties.author = "AI Agents Course"
    document.add_heading(title, level=1)

    for block in content.split("\n\n"):
        document.add_paragraph(block.strip())

    document.save(path)


def find_font_file() -> str | None:
    candidates = [
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"),
    ]

    for candidate in candidates:
        if candidate.exists():
            return str(candidate)

    return None


def write_pdf_file(path: Path, content: str) -> None:
    pdf = fitz.open()
    page = pdf.new_page(width=595, height=842)

    rect = fitz.Rect(50, 50, 545, 792)
    font_file = find_font_file()

    if font_file is not None:
        page.insert_textbox(
            rect,
            content,
            fontsize=10,
            fontfile=font_file,
        )
    else:
        page.insert_textbox(
            rect,
            content,
            fontsize=10,
        )

    pdf.save(path)
    pdf.close()


def write_document(path: Path, title: str, content: str) -> None:
    extension = path.suffix.lower()

    if extension in {".md", ".txt"}:
        write_text_file(path, content)
        return

    if extension == ".html":
        write_html_file(path, title, content)
        return

    if extension == ".docx":
        write_docx_file(path, title, content)
        return

    if extension == ".pdf":
        write_pdf_file(path, content)
        return

    raise ValueError(f"Неподдерживаемое расширение: {path}")


def generate_documents(data_dir: Path, force: bool = False) -> list[Path]:
    created_files = []
    extensions = [".md", ".txt", ".html", ".docx", ".pdf"]

    for category, titles in CATEGORIES.items():
        category_dir = data_dir / category
        category_dir.mkdir(parents=True, exist_ok=True)

        for index, title in enumerate(titles, start=1):
            extension = extensions[(index - 1) % len(extensions)]
            file_name = (
                f"{category}_{index:02d}_v1.{index % 3}"
                f"{extension}"
            )
            file_path = category_dir / file_name

            if file_path.exists() and not force:
                continue

            content = build_content(
                category=category,
                title=title,
                index=index,
            )

            write_document(
                path=file_path,
                title=title,
                content=content,
            )

            created_files.append(file_path)

    return created_files


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Подготовка учебного корпоративного корпуса для RAG.",
    )

    parser.add_argument(
        "--data-dir",
        default="data",
        help="Папка, куда будут записаны документы.",
    )

    parser.add_argument(
        "--force",
        action="store_true",
        help="Перезаписать уже созданные документы.",
    )

    return parser.parse_args()


def main() -> None:
    args = parse_args()
    data_dir = Path(args.data_dir)

    created_files = generate_documents(
        data_dir=data_dir,
        force=args.force,
    )

    print(f"Папка корпуса: {data_dir}")
    print(f"Создано новых файлов: {len(created_files)}")

    if created_files:
        print("Примеры созданных файлов:")
        for file_path in created_files[:10]:
            print(f"- {file_path}")


if __name__ == "__main__":
    main()