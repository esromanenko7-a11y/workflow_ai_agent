import base64
from io import BytesIO

from fastapi import UploadFile
from openai import AsyncOpenAI
from pypdf import PdfReader
from docx import Document


MAX_EXTRACTED_TEXT_CHARS = 30_000
MAX_PDF_PAGES = 50


def extract_pdf_text(data: bytes, max_pages: int = MAX_PDF_PAGES) -> str:
    reader = PdfReader(BytesIO(data))

    pages_text: list[str] = []

    for page in reader.pages[:max_pages]:
        text = page.extract_text() or ""
        if text.strip():
            pages_text.append(text.strip())

    extracted = "\n\n".join(pages_text).strip()

    if len(reader.pages) >= 5 and len(extracted) < 100:
        return (
            "[PDF похож на скан: текст почти не извлечён. "
            "Нужен OCR или текстовая версия документа.]"
        )

    return extracted


def extract_docx_text(data: bytes) -> str:
    document = Document(BytesIO(data))

    chunks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]
            if cells:
                chunks.append(" | ".join(cells))

    return "\n".join(chunks).strip()


async def whisper_transcribe(
    audio_bytes: bytes,
    filename: str,
    client: AsyncOpenAI | None = None,
) -> str:
    """
    Whisper-1 принимает ogg/m4a/mp3/wav/flac/webm напрямую.

    Важно: FFmpeg здесь не нужен.
    """
    openai_client = client or AsyncOpenAI()

    file_obj = BytesIO(audio_bytes)
    file_obj.name = filename

    result = await openai_client.audio.transcriptions.create(
        model="whisper-1",
        file=file_obj,
    )

    return result.text


async def media_to_part(
    media: UploadFile,
    client: AsyncOpenAI | None = None,
) -> dict:
    mime = media.content_type or ""
    data = await media.read()

    if mime.startswith("image/"):
        encoded = base64.b64encode(data).decode("ascii")

        return {
            "type": "image_url",
            "image_url": {
                "url": f"data:{mime};base64,{encoded}",
            },
        }

    if mime.startswith("audio/") or mime == "application/ogg":
        transcript = await whisper_transcribe(
            audio_bytes=data,
            filename=media.filename or "audio.ogg",
            client=client,
        )

        return {
            "type": "text",
            "text": f"[пользователь сказал голосом]:\n{transcript}",
        }

    if mime == "application/pdf":
        text = extract_pdf_text(data)[:MAX_EXTRACTED_TEXT_CHARS]

        return {
            "type": "text",
            "text": f"[документ PDF]:\n{text}",
        }

    if mime.endswith("wordprocessingml.document"):
        text = extract_docx_text(data)[:MAX_EXTRACTED_TEXT_CHARS]

        return {
            "type": "text",
            "text": f"[документ DOCX]:\n{text}",
        }

    raise ValueError(f"Unsupported media type: {mime}")
